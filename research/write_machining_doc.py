#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
write_machining_doc.py — 把加工仿真流程写成 Word 文档，放在 Simple-gr 下。

用法:
  python write_machining_doc.py [out.docx]
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))
import compute_machining as cm


def main():
    out = sys.argv[1] if len(sys.argv) > 1 else \
        str(Path(__file__).resolve().parent.parent / "加工仿真说明.docx")

    mcfg = cm.load_config()
    feed = mcfg.get("feed", 500.0)
    tool_r = mcfg.get("tool_r", 5.0)
    ball_r = mcfg.get("ball_r", 5.0)
    scallop = mcfg.get("scallop", 0.1)
    overhead = mcfg.get("overhead", 4.0)
    point_overhead = mcfg.get("point_overhead", 10.0)

    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 论文格式：正文宋体+Times New Roman 小四(12pt)，标题黑体
    def set_style_font(style, latin, east, size=None, bold=None):
        style.font.name = latin
        if size is not None:
            style.font.size = Pt(size)
        if bold is not None:
            style.font.bold = bold
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn('w:eastAsia'), east)

    set_style_font(doc.styles['Normal'], 'Times New Roman', '宋体', 12)
    for hname, hsize in [('Title', 16), ('Heading 1', 14), ('Heading 2', 12)]:
        try:
            set_style_font(doc.styles[hname], 'Times New Roman', '黑体', hsize, True)
        except Exception:
            pass

    doc.add_heading("直纹面拟合叶片加工仿真说明", level=0)

    doc.add_heading("1. 概述", level=1)
    doc.add_paragraph(
        "本文档描述基于直纹面分片拟合的叶片加工仿真流程。"
        "核心思想：将叶片自由曲面拟合成一组直纹面片（可展面），"
        "从而可用锥度刀侧刃铣（侧铣）一次走刀加工，替代传统球头刀点铣，"
        "在相同精度下比较加工时间并评估提速比。")

    doc.add_heading("2. 工作流程", level=1)
    for s in [
        "① 直纹面拟合：simple.exe --mode ruled 将叶片拟合成直纹面片，导出准线参数 *_params.txt 与网格。",
        "② 刀轨计算：compute_machining.py 读准线，生成侧铣/点铣的 CL 刀位数据（刀心点 + 刀轴方向）。",
        "③ 时间对比：按相同精度指标，比较侧铣（策略 A）与点铣（策略 B）的加工时间，计算提速比。",
    ]:
        doc.add_paragraph(s, style="List Number")

    doc.add_heading("3. 侧铣模型（策略 A）", level=1)
    for s in [
        f"刀具：锥度立铣刀（小锥角退化为圆柱刀），半径 R_t = {tool_r} mm。",
        "刀轴：沿直纹面母线方向（母线 = 上下两准线之差 C1−C0）。",
        "刀心（CL 点）：母线中点沿曲面法向偏置 R_t。",
        "进给：沿准线方向，每片一次走刀。",
        "时间：T_flank = 准线长 / 进给率。",
        "可展性：母线法向扭转角 β 仅作报告统计；轻微不可展由锥度刀锥角补偿。",
    ]:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading("4. 点铣模型（策略 B，传统基线）", level=1)
    for s in [
        f"刀具：球头刀，半径 R_b = {ball_r} mm。",
        "刀心（CL 点）：曲面点沿法向偏置 R_b。",
        f"行距：由残留高度反算 stepover = 2√(2·R·h − h²)，残留 h = {scallop} mm。",
        "面积：使用原始 NURBS 采样网格（blade*_mesh.obj）的面积。",
        "时间：T_point = 原曲面面积 / 行距 / 进给率。",
    ]:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading("5. 时间模型", level=1)
    for s in [
        f"切削时间 = 刀轨长 / 进给率（进给 = {feed} mm/min）。",
        f"侧铣非切削 = 连通域数 × 进退刀开销（{overhead} s/区域，刀轨按连通域蛇形拼接）。",
        f"点铣非切削 = {point_overhead} s（整体一次进退刀）。",
        "提速比 = T_point / T_flank。",
    ]:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading("6. 精度对齐", level=1)
    doc.add_paragraph(
        "侧铣误差 = 拟合误差（直纹面与原曲面的距离），由拟合容差控制（约 0.1 mm）。\n"
        f"点铣误差 = 残留高度（scallop = {scallop} mm）。\n"
        "二者在同一量级（约 0.1 mm）下比较，保证时间对比有意义。")

    doc.add_heading("7. 结果", level=1)
    doc.add_paragraph(
        "示例叶片（blade.igs）：叶盆 5×2 + 叶背 2×6 = 22 片，"
        "拟合 max 误差约 0.09 mm，侧铣相对点铣提速约 3.8 倍。")

    doc.add_heading("8. 假设与局限", level=1)
    for s in [
        "不考虑相邻叶片、夹具的碰撞与刀轴可达性/奇点。",
        "理想机床，进给率恒定，无加减速、无刀轴旋转限速。",
        "点铣行距公式为凸/平面近似，凹面处实际残留偏大。",
        "侧铣处理全部片，轻微不可展由锥度刀补偿。",
    ]:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading("9. 算法输入与输出", level=1)
    for s in [
        "输入：文件 *_params.txt，数学格式：B 样条曲线的定义参数。",
        "输出：文件 toolpath_*.csv、toolpath_*.vtk、toolpath_*.dxf。",
        "备注：csv 文件用于绘制折线，vtk 文件用于 ParaView 可视化，dxf 文件可直接导入 NX。",
    ]:
        doc.add_paragraph(s, style="List Bullet")

    doc.save(out)
    print(f"[doc] -> {out}")


if __name__ == "__main__":
    main()
